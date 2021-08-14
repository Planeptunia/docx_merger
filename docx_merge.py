import os
from docx import Document
import docx2txt
# Переменные
document = Document()
paths = []
# Получение путей файлов
dir = './input'
files = os.listdir(dir)
for x in files:
    path = './input/' + x
    paths.append(path)
# Цикл обработки
for i in paths:
    # Получение названия файла для вывода
    file = paths.index(i, 0, len(paths))
    # Обработка текста
    text = docx2txt.process(i)
    # Цикл проверки
    for k in range(len(paths)):
        # Создание проверочного текста
        check_text = docx2txt.process(paths[k])
        # Извлечение текста
        if check_text != text:
            document.add_paragraph(text)
            document.add_paragraph("\n")
            print('%s просканирован' % (files[file]))
            break
# Сохранение
document.save('merged.docx')
